import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    """Set the background color of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set custom padding (margins) for a specific cell in dxa (1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="B0B0B0", sz="4", val="single"):
    """Set clean borders for the table."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def main():
    doc = Document()

    # 1. Page Setup: A4 with 2.5 cm margins all around (standard clean layout)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Calculate printable width: 21 - 2.5*2 = 16.0 cm
    printable_width_cm = 16.0

    # Set default fonts
    # We want "微軟正黑體" (Microsoft JhengHei) as primary font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微軟正黑體'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Dark gray body text

    # Set East Asian font in XML
    rPr = style.element.rPr if style.element.rPr is not None else OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(rFonts)

    # 2. Spacing / Title Layout
    # Top spacing: add some empty paragraphs or spacing before
    title_spacing = doc.add_paragraph()
    title_spacing.paragraph_format.space_before = Pt(40)
    title_spacing.paragraph_format.space_after = Pt(20)

    # Title: 三年五班作文
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('三年五班作文')
    title_run.font.name = '微軟正黑體'
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Premium Navy Blue

    # Subtitle: 作文檢閱封面
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run('— 學生作文檢閱紀錄表 —')
    subtitle_run.font.name = '微軟正黑體'
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D) # Cool Grey
    subtitle_p.paragraph_format.space_after = Pt(60) # Space before table

    # 3. Composition Topics Table
    # 2 columns: Title (11 cm), Score (5 cm) = Total 16 cm (matches printable area)
    col_widths = [Cm(11.0), Cm(5.0)]
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="CCCCCC", sz="6") # Light gray borders, size 6 is 3/4 pt

    # Headers
    headers = ["作文題目", "登記分數"]
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = col_widths[i]
        set_cell_background(cell, "EAECEE") # Light gray-blue header
        set_cell_margins(cell, top=160, bottom=160, left=200, right=200) # Comfortable margins
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Row Heights and Content
    table.rows[0].height = Cm(1.2) # Header height
    
    topics = [
        "1. 我的好朋友",
        "2. 我的媽媽",
        "3. 飛行員與小王子（改寫故事）"
    ]

    for idx, topic in enumerate(topics):
        row = table.rows[idx + 1]
        row.height = Cm(2.2) # Tall height for writing score/comments
        
        # Col 0: Title
        cell_title = row.cells[0]
        cell_title.width = col_widths[0]
        set_cell_margins(cell_title, top=200, bottom=200, left=300, right=200)
        p_title = cell_title.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_title = p_title.add_run(topic)
        run_title.font.size = Pt(14)
        run_title.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        cell_title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Col 1: Score Space
        cell_score = row.cells[1]
        cell_score.width = col_widths[1]
        set_cell_margins(cell_score, top=200, bottom=200, left=200, right=200)
        p_score = cell_score.paragraphs[0]
        p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add a subtle placeholder or leave empty for teacher writing
        run_score = p_score.add_run("") 
        cell_score.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add spacing after table
    post_table_spacing = doc.add_paragraph()
    post_table_spacing.paragraph_format.space_before = Pt(80)

    # 4. Student Information Section
    # We will use a borderless table to align "班級：三年五班", "座號：__________", "姓名：__________" nicely.
    info_table = doc.add_table(rows=1, cols=3)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info_widths = [Cm(5.0), Cm(5.0), Cm(6.0)]
    info_cells = info_table.rows[0].cells
    
    info_fields = [
        ("班級：三年五班", WD_ALIGN_PARAGRAPH.LEFT),
        ("座號：__________", WD_ALIGN_PARAGRAPH.CENTER),
        ("姓名：_______________", WD_ALIGN_PARAGRAPH.RIGHT)
    ]
    
    for i, (text, align) in enumerate(info_fields):
        cell = info_cells[i]
        cell.width = info_widths[i]
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        run.font.name = '微軟正黑體'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Save the document
    output_filename = "三年五班作文檢閱封面.docx"
    doc.save(output_filename)
    print(f"Document successfully created and saved to {output_filename}")

if __name__ == "__main__":
    main()
