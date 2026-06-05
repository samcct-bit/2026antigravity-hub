#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
格式匯出工具：將雙語對照 JSON 轉換為雙欄 Excel 表格與 Word 表格。
"""
import argparse
import json
import os
import sys

def export_xlsx(data, notes, output_path):
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "雙語對照翻譯"
    
    # 樣式設定
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(name="Microsoft JhengHei", size=11, bold=True, color="FFFFFF")
    cell_font = Font(name="Microsoft JhengHei", size=10)
    note_font = Font(name="Microsoft JhengHei", size=9, italic=True, color="595959")
    
    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )
    
    # 標題
    ws.append(["原文", "指定語言譯文"])
    for col in [1, 2]:
        cell = ws.cell(row=1, column=col)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    
    # 內容
    for item in data:
        orig = item.get("original", "")
        trans = item.get("translation", "")
        ws.append([orig, trans])
        
        row_idx = ws.max_row
        for col in [1, 2]:
            cell = ws.cell(row=row_idx, column=col)
            cell.font = cell_font
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            
    # 註釋/補充
    if notes:
        ws.append([]) # 空行
        row_idx = ws.max_row
        ws.cell(row=row_idx, column=1, value="【翻譯註釋與補充說明】").font = Font(name="Microsoft JhengHei", size=10, bold=True)
        for line in notes.split("\n"):
            if line.strip():
                ws.append([line.strip(), ""])
                ws.cell(row=ws.max_row, column=1).font = note_font
                # 合併儲存格
                ws.merge_cells(start_row=ws.max_row, start_column=1, end_row=ws.max_row, end_column=2)

    # 寬度調整
    ws.column_dimensions['A'].width = 50
    ws.column_dimensions['B'].width = 50
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"Excel 匯出成功：{output_path}")

def export_docx(data, notes, output_path):
    import docx
    from docx.shared import Inches, Pt
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
    
    doc = docx.Document()
    
    # 設定字型
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    
    doc.add_heading('雙語對照翻譯報告', level=1)
    
    # 建立表格
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # 表頭設定
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '原文'
    hdr_cells[1].text = '指定語言譯文'
    
    # 設定表頭背景色 (深藍)
    shading_elm1 = parse_xml(r'<w:shd {} w:fill="1F4E78"/>'.format(nsdecls('w')))
    shading_elm2 = parse_xml(r'<w:shd {} w:fill="1F4E78"/>'.format(nsdecls('w')))
    hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm1)
    hdr_cells[1]._tc.get_or_add_tcPr().append(shading_elm2)
    
    # 表頭文字白色與粗體
    for cell in hdr_cells:
        for p in cell.paragraphs:
            p.alignment = 1 # Center
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = docx.shared.RGBColor(255, 255, 255)
                
    # 填充表格內容
    for item in data:
        orig = item.get("original", "")
        trans = item.get("translation", "")
        row = table.add_row()
        row.cells[0].text = orig
        row.cells[1].text = trans
        
    # 表格寬度調整
    for row in table.rows:
        row.cells[0].width = Inches(3.2)
        row.cells[1].width = Inches(3.2)
        
    # 註釋
    if notes:
        doc.add_paragraph() # 空行
        p_title = doc.add_paragraph()
        r_title = p_title.add_run("【翻譯註釋與補充說明】")
        r_title.bold = True
        r_title.font.size = Pt(11)
        
        for line in notes.split("\n"):
            if line.strip():
                p_note = doc.add_paragraph()
                r_note = p_note.add_run(line.strip())
                r_note.italic = True
                r_note.font.size = Pt(9.5)
                r_note.font.color.rgb = docx.shared.RGBColor(89, 89, 89)
                
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Word 匯出成功：{output_path}")

def export_md(data, notes, output_path):
    lines = []
    lines.append("# 雙語對照翻譯對照表\n")
    lines.append("| 原文 | 指定語言譯文 |")
    lines.append("| :--- | :--- |")
    for item in data:
        orig = item.get("original", "").replace("\n", "<br>")
        trans = item.get("translation", "").replace("\n", "<br>")
        lines.append(f"| {orig} | {trans} |")
    
    if notes:
        lines.append("\n### 註釋與說明\n")
        lines.append(notes)
        
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Markdown 匯出成功：{output_path}")

def main():
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
        
    parser = argparse.ArgumentParser(description="匯出對照格式檔案")
    parser.add_argument("--json", required=True, help="翻譯對照 JSON 檔路徑")
    parser.add_argument("--out_xlsx", help="輸出 Excel 檔路徑")
    parser.add_argument("--out_docx", help="輸出 Word 檔路徑")
    parser.add_argument("--out_md", help="輸出 Markdown 檔路徑")
    parser.add_argument("--notes", default="", help="翻譯註釋說明文字")
    args = parser.parse_args()
    
    with open(args.json, "r", encoding="utf-8") as f:
        data = json.load(f)
        
    # 如果 data 是個 dict，試圖取出內含的列表
    if isinstance(data, dict):
        data = data.get("translations", data.get("gems", []))
        
    if args.out_xlsx:
        export_xlsx(data, args.notes, args.out_xlsx)
    if args.out_docx:
        export_docx(data, args.notes, args.out_docx)
    if args.out_md:
        export_md(data, args.notes, args.out_md)

if __name__ == "__main__":
    main()
